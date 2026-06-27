from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "quant_model_logic_report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(95, 108, 124)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
RISK_RED = RGBColor(155, 28, 28)
GOLD = RGBColor(122, 90, 0)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    # Named override for Chinese text rendering in Word/LibreOffice.
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def paragraph_border_bottom(paragraph, color="A8B5C6", size="8", space="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_para(doc, text="", style=None, size=None, color=None, bold=None, italic=None, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    colors = {1: BLUE, 2: BLUE, 3: DARK_BLUE}
    sizes = {1: 16, 2: 13, 3: 12}
    set_run_font(run, size=sizes.get(level, 12), color=colors.get(level, BLUE), bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) and node.get(qn("w:abstractNumId")).isdigit()
    ]
    existing_num = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) and node.get(qn("w:numId")).isdigit()
    ]
    abstract_id = (max(existing_abs) + 1) if existing_abs else 1
    num_id = (max(existing_num) + 1) if existing_num else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered(doc, items):
    num_id = add_numbering_definition(doc)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_pr.append(ilvl)
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_node)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_callout(doc, title, body, kind="info"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E0" if kind == "warn" else LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=GOLD if kind == "warn" else DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    add_para(doc, "", after=4)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.18
            run = p.add_run(str(value))
            set_run_font(run, size=9.2)
    set_table_geometry(table, widths)
    add_para(doc, "", after=4)
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("Global Quant Watch")
    set_run_font(left, size=9, color=MUTED, bold=True)
    header.add_run("  |  Model Logic & Agent Training")
    set_run_font(header.runs[-1], size=9, color=MUTED)
    paragraph_border_bottom(header, color="D5DCE6", size="4", space="8")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Decision support only - not investment advice")
    set_run_font(run, size=8.5, color=MUTED)


def build_doc():
    doc = Document()
    setup_styles(doc)

    add_para(doc, "GLOBAL QUANT WATCH", size=10, color=MUTED, bold=True, after=4, before=8)
    title = add_para(doc, "量化预测模型与实盘训练 Agent 说明书", size=24, color=INK, bold=True, after=6)
    title.paragraph_format.line_spacing = 1.05
    add_para(
        doc,
        f"版本：{date.today().isoformat()}  |  适用市场：ASX / US / CN  |  文档用途：解释预测逻辑、训练机制、过拟合控制与下一步提升路线",
        size=10.5,
        color=MUTED,
        after=10,
    )
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="A8B5C6", size="8", space="10")

    add_callout(
        doc,
        "核心结论",
        "这次升级的重点不是机械提高置信度，而是把置信度做成更接近“可校准概率”的东西：只有当真实行情、新闻/宏观、因子、历史相似路径、模型集成和 agent 训练同时支持时，才允许置信度上行。新增三个实盘训练 agent 后，系统会以真实市场数据做纸面交易、记录止损和失败经验，并用长期记忆微调策略阈值。",
    )
    add_callout(
        doc,
        "风险边界",
        "当前系统是交易决策辅助和纸面实盘训练系统，不直接替你向券商下真实订单。所有预测都可能失败；市场存在非平稳、突发事件、数据延迟、流动性和滑点问题。高置信度不等于保证收益。",
        kind="warn",
    )

    add_heading(doc, "1. 本次已完成的模型升级", 1)
    add_bullets(
        doc,
        [
            "训练层从 2 个 agent 扩展到 5 个：趋势、回撤、突破高频、新闻资金流、稳健配置。",
            "候选策略从 4 类扩展到 10 类，覆盖趋势回踩、放量突破、超跌反弹、低波动突破、新闻确认趋势、利好修复和现金保护入场。",
            "新增 evidence quality calibration：把新闻数量、真实因子覆盖、模型共识、历史相似路径、数据源降级和近期买入命中率一起用于置信度校准。",
            "每个 agent 的买入阈值、止损、止盈、最长持有期和仓位大小按照风格独立设置，避免所有策略同质化。",
            "止损或亏损退出会写入亏损复盘 lesson，长期记忆继承这些经验，用于降低同类形态或同一股票的偏置。",
        ]
    )

    add_table(
        doc,
        ["Agent", "角色", "交易频率倾向", "入场核心", "风控逻辑"],
        [
            ("趋势 Agent", "捕捉中短期动量延续", "中等", "趋势分、MACD、目标涨幅、量能", "常规止损，盈利后比较是否过早止盈"),
            ("回撤 Agent", "寻找超跌后的修复", "中等偏低", "RSI 低位、最大上探空间、风险回落", "更紧止损，失败后减少反弹形态频率"),
            ("突破实盘训练 Agent", "提高试错和交易频率", "较高", "放量、20日高位、5日走势不弱", "小仓位、快止损、短持有"),
            ("新闻资金流 Agent", "把消息面转成交易证据", "中等", "新闻/因子证据、资金流、趋势确认", "新闻证据薄时禁止升置信"),
            ("稳健配置 Agent", "控制回撤和组合暴露", "低", "风险分、流动性、共识和现金保护", "单票小仓位，保留补仓/对冲空间"),
        ],
        [1300, 2050, 1050, 2700, 2260],
    )

    add_heading(doc, "2. 一次完整预测如何发生", 1)
    add_numbered(
        doc,
        [
            "读取市场数据：按市场和标的读取真实 OHLCV、实时报价、指数状态、成交/订单流可用数据，并优先使用本地真实缓存避免休市重复消耗额度。",
            "清洗与校验：统一代码、市场、时间、价格和成交量；拒绝用 ETF 冒充现金指数；数据源降级时保留 warning 并压低置信。",
            "构造技术特征：计算 RSI、MACD、均线、波动率、5日变化、成交量比、趋势分、动量分、风险分和历史相似路径。",
            "读取基本面与外部因子：公告、财报、宏观、板块、资金/空头、流动性、市场状态、相对强弱、回测校准等。",
            "读取新闻与全球事件：股票直接新闻、行业上下游、竞品、宏观、央行、地缘政治、政策官员讲话、X/YouTube 等一手/准一手信号。",
            "AI 和本地模型生成初始判断：规则模型、历史相似模型、因子模型、AI 文本推理和市场 regime 层共同生成预估涨幅、最大上探、下跌风险和动作建议。",
            "置信度校准：根据近期命中率、买入达标率、证据质量、模型一致性和数据降级情况调整置信度，不把弱数据包装成高置信。",
            "agent 覆盖层：5 个 agent 根据各自阈值决定是否纸面买入、加仓、持有、卖出或止损，并把结果写入长期记忆。",
            "组合与仓位建议：结合总资金、可用资金、单票上限、当前持仓、后续补仓空间、风险暴露和对冲属性，给出更保守的交易建议。",
        ]
    )

    add_heading(doc, "3. 新闻解读层：从文本到交易因子", 1)
    add_para(
        doc,
        "新闻不是只看“公司名是否出现”。系统应把新闻拆成直接、行业、上下游、竞品、宏观、利率、政策、战争/地缘政治和大盘风险几个通道。每条新闻先提取实体、事件、方向、时间敏感度、影响范围和可信度，再映射到标的可能受益或受损的路径。",
    )
    add_table(
        doc,
        ["新闻通道", "提取内容", "如何影响预测", "防错处理"],
        [
            ("公司直接", "财报、公告、订单、管理层、监管", "直接改变盈利预期和短期情绪", "核对日期和来源，避免旧闻当新信号"),
            ("产业链", "上游成本、下游需求、运输、能源、原材料", "改变毛利率、销量和交付确定性", "区分一次性事件和持续趋势"),
            ("竞品/替代", "竞品涨价、事故、产能、政策倾斜", "可能形成替代受益或竞争压力", "需要行业关系图，不能只靠关键词"),
            ("宏观/利率", "央行、通胀、就业、汇率、债券收益率", "影响估值折现率、风险偏好和资金流", "按市场和行业敏感度加权"),
            ("地缘政治", "战争、制裁、贸易限制、能源冲击", "可能压制大盘或推升防御/资源类资产", "不默认所有战争都让所有股票下跌，按行业暴露分解"),
            ("社交/视频", "X 热点、关键人物发言、YouTube 热搜", "用于捕捉早期情绪和传播速度", "降低权重，要求与价格/成交/可信来源交叉验证"),
        ],
        [1450, 2250, 2600, 3060],
    )

    add_heading(doc, "4. 技术面与因子层", 1)
    add_bullets(
        doc,
        [
            "技术面：RSI 衡量短期超买超卖，MACD 衡量动量变化，均线和趋势分衡量方向，量比衡量资金参与度，波动率和回撤衡量风险。",
            "成交与订单流：在有真实 tick/L1/L2 授权数据时，进一步识别主动买入、主动卖出、被动成交、价格阶梯和冲击成本；没有真实数据时不伪造成逐笔。",
            "历史相似路径：寻找过去走势、成交量、波动和价格结构相似的片段，统计后续最大上探、最终收益、先止损概率和达到目标概率。",
            "基本面因子：公告/财报、盈利质量、估值、行业景气、宏观敏感度、资金流、相对强弱、流动性和校准表现。",
            "因子合格门槛：无量纲、丰富度、无未来函数、缺失值处理、极端值处理、标准化六项必须通过，否则降低权重。",
        ]
    )
    add_table(
        doc,
        ["因子质量检查", "标准", "失败时的处理"],
        [
            ("无量纲", "不同价格/市值/币种股票可比较", "做收益率、分位数、z-score 或行业中性化"),
            ("丰富度", "信息维度互补，不重复堆相同指标", "删除高度相关或边际贡献低的重复因子"),
            ("无未来函数", "决策时点必须已经可见", "使用 point-in-time 数据和公告发布日期"),
            ("缺失值", "缺失原因可解释", "单独标记缺失，不把缺失直接当 0"),
            ("极端值", "异常值不支配模型", "winsorize、robust scaling 或异常剔除"),
            ("标准化", "训练/预测同一变换", "只用训练窗口拟合 scaler，避免泄漏"),
        ],
        [1700, 4550, 3110],
    )

    add_heading(doc, "5. 当前模型集合", 1)
    add_table(
        doc,
        ["模型/模块", "输入", "输出", "主要用途"],
        [
            ("规则技术模型", "OHLCV、RSI、MACD、均线、量比", "趋势/动量/风险分", "基础方向判断和入场过滤"),
            ("历史相似路径", "走势、成交、波动、价格结构", "未来收益分布、止损先触发概率", "让预测参考过去类似市场经验"),
            ("因子评分模型", "基本面、宏观、资金、流动性、相对强弱", "factor score 和证据覆盖", "避免只看日内涨跌"),
            ("新闻文本推理", "多源新闻、社交、视频、宏观事件", "利好/利空、影响路径、强弱相关", "把消息面转成可解释证据"),
            ("集成共识层", "规则、AI、历史、因子、市场状态", "预测涨幅、置信度、动作", "减少单模型偏差"),
            ("校准层", "历史命中率、Brier/命中桶、数据降级", "校准后的置信度", "防止虚高置信和过拟合"),
            ("agent 纸面训练", "实时分析结果和当前价格", "买/卖/持有、亏损 lesson", "用策略表现反推阈值和仓位"),
        ],
        [1550, 2850, 2350, 2610],
    )

    add_heading(doc, "6. 置信度为什么不能硬拉高", 1)
    add_para(
        doc,
        "置信度应代表“在当前证据条件下，目标事件发生的校准概率”。如果为了接近 75% 目标而简单加分，短期页面会好看，但真实交易会更容易过拟合。正确做法是让强证据提高置信，让弱证据和近期失败自动压低置信。",
    )
    add_bullets(
        doc,
        [
            "允许升置信的条件：新闻/因子覆盖足够、模型共识高、历史相似路径支持、近期买入达标率不差、行情源未严重降级。",
            "必须降置信的条件：新闻为空、真实因子太少、单源降级、近期买入命中率低、预测过度依赖当天涨跌。",
            "预测涨幅也要校准：弱证据时不只降置信，还要缩小预估涨幅和最大上探空间。",
            "长期准确率提升来自记录每次预测样本、延迟回看结果、按时间切分训练/验证，而不是当天涨跌一变就马上反向。",
        ]
    )

    add_heading(doc, "7. 实盘训练 agent 如何提高策略", 1)
    add_para(
        doc,
        "这里的“实盘训练”指用真实市场数据做纸面交易，不直接下真实订单。每个 agent 在刷新时按照自己的风格评估候选股票：满足阈值就纸面买入，触及止损、目标、周期或信号转弱就卖出，并把结果写入记忆库。",
    )
    add_numbered(
        doc,
        [
            "历史回放：用最近约 220 根 K 线测试 10 类候选策略，统计交易次数、胜率、平均收益和最大回撤。",
            "策略评分：score = 平均收益、胜率偏离、交易样本数和回撤惩罚的组合；交易太少不会被直接认定为最优。",
            "阈值自调：表现好的策略降低一点入场阈值，提高交易机会；表现差的策略提高阈值或降低个股偏置。",
            "仓位自控：突破 agent 小仓高频，稳健 agent 小仓低频，趋势 agent 中等仓位，避免有现金就全买一只。",
            "失败学习：止损、信号转弱、临近周期仍亏损等原因会形成 lesson，并在后续训练中降低同类形态的信心。",
        ]
    )

    doc.add_page_break()
    add_heading(doc, "8. 如何继续提高准确率", 1)
    add_table(
        doc,
        ["优先级", "改进方向", "为什么有效", "注意事项"],
        [
            ("1", "更干净的 point-in-time 数据", "避免未来函数，是所有模型准确率的地基", "财报、公告、成分股、复权必须按当时可见状态"),
            ("2", "预测样本本地累计", "让校准层看到足够多失败和成功样本", "按市场、行业、周期分桶，不能混成一个大平均"),
            ("3", "更强标签体系", "区分最终收益、最大上探、先止损、触达目标", "15日涨5%和次日方向是不同任务"),
            ("4", "LightGBM/线性基线", "对中小样本、非线性因子通常更稳", "严格时间序列切分和早停"),
            ("5", "LSTM/Transformer", "可学习序列状态和 regime", "需要更多数据，否则比树模型更容易过拟合"),
            ("6", "交易成本/滑点模型", "让纸面收益更接近真实收益", "低流动性股票必须惩罚"),
            ("7", "组合优化", "提高整体收益/回撤，而不是单票命中", "使用行业、因子、市场 beta 和现金约束"),
        ],
        [800, 2400, 3100, 3060],
    )

    add_heading(doc, "9. 借鉴的开源量化项目", 1)
    add_table(
        doc,
        ["项目", "可借鉴点", "对本系统的启发"],
        [
            ("Microsoft Qlib", "覆盖监督学习、市场动态建模、强化学习和从 alpha 到订单执行的研究管线", "把因子、模型训练、回测、风险和执行拆成清晰模块；后续可接 LightGBM/LSTM/Transformer"),
            ("FinRL", "以市场环境、DRL agent 和金融应用三层组织，并常用 A2C/DDPG/PPO/TD3/SAC 等 agent", "强化学习适合做仓位和交易时机实验，但必须先有稳定回测和风控"),
            ("QuantConnect LEAN", "事件驱动、模块化、支持回测/优化/实盘和多类数据", "长期应向事件驱动架构靠拢，便于统一指数、个股、新闻和订单事件"),
            ("Freqtrade", "提供 dry-run、SQLite 持久化、回测、画图和机器学习策略优化", "纸面训练、持久化日志、策略迭代和风险提示是实盘前必要环节"),
        ],
        [1700, 3900, 3760],
    )
    add_para(
        doc,
        "来源：Qlib https://github.com/microsoft/qlib；FinRL https://github.com/AI4Finance-Foundation/FinRL；QuantConnect LEAN https://github.com/QuantConnect/Lean；Freqtrade https://github.com/freqtrade/freqtrade",
        size=8.8,
        color=MUTED,
        after=6,
    )

    add_heading(doc, "10. 市场是否本身难以预测", 1)
    add_para(
        doc,
        "是的，特别是短周期。股票短期收益的信噪比很低，市场结构会变化，公开信息会被快速定价，突发事件和流动性冲击会让历史规律失效。因此系统不能承诺长期稳定高胜率，只能持续提高“可证伪、可校准、可复盘”的能力。",
    )
    add_bullets(
        doc,
        [
            "更多数据有帮助，但只有在数据干净、时间戳正确、没有未来函数、覆盖足够长的不同市场 regime 时才有帮助。",
            "更复杂模型不一定更好。小样本下，LightGBM/正则化线性模型/简单集成往往比深度模型更稳。",
            "提高收益不等于只提高方向预测。仓位、止损、卖出纪律、交易成本和组合分散通常同样重要。",
            "系统最应该优化的是校准后的期望收益：概率 × 幅度 - 风险 - 成本，而不是单一置信度。",
        ]
    )

    add_heading(doc, "11. 后续建议路线", 1)
    add_numbered(
        doc,
        [
            "把每次预测、实际 1/5/15 日结果、是否触达目标、是否先触发止损全部持久化到本地 record。",
            "按市场和行业建立独立校准桶，避免 ASX、US、A股混在一起训练。",
            "接入 Qlib 数据管线和 LightGBM 基线，先验证因子有效性，再接 LSTM/Transformer。",
            "建立特征重要性和漂移检测：某个因子连续失效时自动降权，而不是继续硬用。",
            "把 agent 的纸面交易收益改为含交易成本、滑点、最小成交量和最大持仓限制的净收益。",
            "建立 walk-forward 回测：训练窗口、验证窗口、上线窗口分离，所有参数调整必须在未来窗口验证。",
            "对新闻文本模型加入事件类型和行业知识图谱，让强相关、弱相关、传导相关分开计权。",
        ]
    )

    doc.add_page_break()
    add_heading(doc, "12. 当前实现的可解释公式", 1)
    add_para(
        doc,
        "系统并不是单一公式给出买卖，而是多层打分。简化后可以理解为：",
    )
    add_callout(
        doc,
        "决策分数",
        "Expected Edge = P(calibrated) x Expected Upside - Downside Risk - Data Penalty - Cost/Slippage + Agent Memory Bias。\n只有当 Expected Edge、策略约束、仓位约束和止损约束同时通过时，才推荐买入或进入观察。",
    )
    add_para(
        doc,
        "置信度上调来自证据质量和历史校准，不来自人为目标。比如你设置“15 日内上涨 5%、置信 75%”，系统会筛选并提醒接近该条件的股票；如果当前证据只能支持 52%，它应如实显示 52%，而不是为了满足目标显示 75%。",
    )

    add_heading(doc, "13. 交付说明", 1)
    add_bullets(
        doc,
        [
            "代码层已加入 3 个新增 agent、10 类候选策略、证据质量校准和亏损 lesson 记录。",
            "这些改动会在刷新分析后逐步积累样本；初期置信度不会突然大幅上升，这是为了避免过拟合。",
            "如果要进一步提升，下一步优先做本地 record 持久化审计、Qlib/LightGBM 训练入口和 walk-forward 评估面板。",
        ]
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
