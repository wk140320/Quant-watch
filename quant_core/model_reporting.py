from __future__ import annotations

import gzip
import hashlib
import html
import json
import math
import os
import random
import re
import statistics
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable

from paper_agents import load_state as load_paper_agent_state


MARKETS = ("ASX", "US", "CN")
MODEL_OUTPUTS = (
    ("directionProbability", "Calibrated final-direction ensemble", "actualDirection"),
    ("ridgeDirectionPrediction", "Logistic direction baseline", "actualDirection"),
    ("treeDirectionPrediction", "Shallow-tree direction challenger", "actualDirection"),
    ("ridgePrediction", "Ridge cross-sectional baseline", "actualTarget"),
    ("elasticPrediction", "ElasticNet probability baseline", "actualTarget"),
    ("lightgbmPrediction", "Shallow-tree path challenger (CatBoost/LightGBM)", "actualTarget"),
    ("pathSafetyPrediction", "Target-before-stop path model", "actualTarget"),
    ("eventPrediction", "Point-in-time event model", "actualTarget"),
    ("ensembleProbability", "Constrained OOF path ensemble", "actualTarget"),
    ("targetProbability", "Calibrated target probability", "actualTarget"),
)
PRODUCTION_THRESHOLDS = {
    "minRowsPerHorizon": 50_000,
    "minOofTestRows": 1_000,
    "minIndependentTestDates": 120,
    "minTargetEvents": 500,
    "minStopEvents": 500,
    "minFolds": 5,
    "minPositiveFolds": 4,
    "minBrierSkill": 0.0,
    "maxEcePct": 5.0,
    "minCalibrationSlope": 0.8,
    "maxCalibrationSlope": 1.2,
    "minProbabilityBucketEvents": 30,
    "maxFeaturePsi": 0.40,
}


def number(value: Any, fallback: float = 0.0) -> float:
    try:
        parsed = float(value)
        return parsed if math.isfinite(parsed) else fallback
    except (TypeError, ValueError):
        return fallback


def clamp(value: Any, low: float, high: float) -> float:
    return max(low, min(high, number(value, low)))


def iso_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def stable_hash(value: Any, length: int = 20) -> str:
    body = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str)
    return hashlib.sha256(body.encode("utf-8")).hexdigest()[:length]


def read_json(path: Path, fallback: Any = None) -> Any:
    try:
        return json.loads(path.read_text("utf-8"))
    except (OSError, ValueError, TypeError):
        return fallback


def metric(value: Any, *, reason: str = "") -> dict[str, Any]:
    parsed = None
    try:
        candidate = float(value)
        if math.isfinite(candidate):
            parsed = round(candidate, 6)
    except (TypeError, ValueError):
        pass
    return {
        "value": parsed,
        "available": parsed is not None,
        "reason": "" if parsed is not None else (reason or "Evidence is unavailable."),
    }


def prediction_id(row: dict[str, Any], manifest: dict[str, Any] | None = None) -> str:
    manifest = manifest or {}
    identity = {
        "market": str(row.get("market") or "").upper(),
        "symbol": str(row.get("symbol") or "").upper(),
        "signalAt": row.get("signalAt") or row.get("date"),
        "horizon": int(number(row.get("horizon"), 0)),
        "target": number(row.get("targetPct"), number(row.get("targetUpside"))),
        "stop": number(row.get("stopPct"), number(row.get("stopLoss"))),
        "labelDefinition": manifest.get("label_definition") or row.get("labelDefinition") or "unknown",
        "featureSchema": manifest.get("feature_schema_hash") or row.get("featureSchemaHash") or "unknown",
        "modelVersion": manifest.get("model_version") or row.get("modelVersion") or "unknown",
    }
    return stable_hash(identity, 32)


def newey_west_mean_ci(values: list[float], max_lag: int | None = None) -> dict[str, Any]:
    clean = [float(value) for value in values if math.isfinite(float(value))]
    count = len(clean)
    if count < 8:
        return {"available": False, "reason": "Fewer than eight ordered observations.", "samples": count}
    center = statistics.fmean(clean)
    residuals = [value - center for value in clean]
    lag = max_lag if max_lag is not None else max(1, int(4 * (count / 100) ** (2 / 9)))
    lag = min(max(1, int(lag)), count - 1)
    long_run_variance = sum(value * value for value in residuals) / count
    for offset in range(1, lag + 1):
        covariance = sum(residuals[index] * residuals[index - offset] for index in range(offset, count)) / count
        long_run_variance += 2 * (1 - offset / (lag + 1)) * covariance
    standard_error = math.sqrt(max(0.0, long_run_variance) / count)
    return {
        "available": True,
        "samples": count,
        "lags": lag,
        "estimate": round(center, 6),
        "standardError": round(standard_error, 6),
        "low": round(center - 1.959963984540054 * standard_error, 6),
        "high": round(center + 1.959963984540054 * standard_error, 6),
        "method": "Newey-West HAC 95% CI",
    }


def _classification_counts(actuals: list[int], predictions: list[int]) -> tuple[int, int, int, int]:
    tp = sum(1 for actual, predicted in zip(actuals, predictions) if actual == predicted == 1)
    tn = sum(1 for actual, predicted in zip(actuals, predictions) if actual == predicted == 0)
    fp = sum(1 for actual, predicted in zip(actuals, predictions) if actual == 0 and predicted == 1)
    fn = sum(1 for actual, predicted in zip(actuals, predictions) if actual == 1 and predicted == 0)
    return tp, tn, fp, fn


def _auc(actuals: list[int], scores: list[float]) -> float | None:
    positive_count = sum(1 for actual in actuals if actual == 1)
    negative_count = len(actuals) - positive_count
    if not positive_count or not negative_count:
        return None
    ordered = sorted(zip(scores, actuals), key=lambda row: row[0])
    positive_rank_sum = 0.0
    index = 0
    while index < len(ordered):
        end = index + 1
        while end < len(ordered) and ordered[end][0] == ordered[index][0]:
            end += 1
        average_rank = ((index + 1) + end) / 2.0
        positive_rank_sum += average_rank * sum(actual for _, actual in ordered[index:end])
        index = end
    return (positive_rank_sum - positive_count * (positive_count + 1) / 2.0) / (positive_count * negative_count)


def _average_precision(actuals: list[int], scores: list[float]) -> float | None:
    positive_count = sum(actuals)
    if positive_count == 0:
        return None
    ordered = sorted(zip(scores, actuals), reverse=True)
    hits = 0
    precision_sum = 0.0
    for index, (_, actual) in enumerate(ordered, start=1):
        if actual:
            hits += 1
            precision_sum += hits / index
    return precision_sum / positive_count


def _platt_slope(actuals: list[int], scores: list[float]) -> tuple[float | None, float | None]:
    if len(set(actuals)) < 2 or len(actuals) < 10:
        return None, None
    intercept = 0.0
    slope = 1.0
    logits = [math.log(clamp(score, 0.001, 0.999) / (1 - clamp(score, 0.001, 0.999))) for score in scores]
    for _ in range(30):
        grad_i = 0.0
        grad_s = 0.0
        hessian_ii = 1e-8
        hessian_is = 0.0
        hessian_ss = 1e-8
        for logit, actual in zip(logits, actuals):
            prediction = 1 / (1 + math.exp(-clamp(intercept + slope * logit, -24, 24)))
            error = prediction - actual
            curvature = max(1e-8, prediction * (1 - prediction))
            grad_i += error
            grad_s += error * logit
            hessian_ii += curvature
            hessian_is += curvature * logit
            hessian_ss += curvature * logit * logit
        determinant = hessian_ii * hessian_ss - hessian_is * hessian_is
        if abs(determinant) < 1e-12:
            break
        delta_i = (hessian_ss * grad_i - hessian_is * grad_s) / determinant
        delta_s = (-hessian_is * grad_i + hessian_ii * grad_s) / determinant
        intercept -= clamp(delta_i, -2.0, 2.0)
        slope -= clamp(delta_s, -1.0, 1.0)
        slope = clamp(slope, 0.02, 5.0)
        if abs(delta_i) < 1e-7 and abs(delta_s) < 1e-7:
            break
    return slope, intercept


def classification_metrics(rows: list[dict[str, Any]], probability_key: str, actual_key: str = "actualTarget") -> dict[str, Any]:
    valid = [
        row for row in rows
        if row.get(probability_key) is not None and row.get(actual_key) is not None
    ]
    if not valid:
        return {"available": False, "reason": "No strict OOF probability rows.", "samples": 0}
    actuals = [1 if number(row.get(actual_key)) >= 0.5 else 0 for row in valid]
    scores = [clamp(row.get(probability_key), 0.001, 0.999) for row in valid]
    predicted = [1 if score >= 0.5 else 0 for score in scores]
    tp, tn, fp, fn = _classification_counts(actuals, predicted)
    positive_support = tp + fn
    negative_support = tn + fp
    precision = tp / max(1, tp + fp)
    recall = tp / max(1, positive_support)
    f1 = 2 * precision * recall / max(1e-12, precision + recall)
    negative_precision = tn / max(1, tn + fn)
    negative_recall = tn / max(1, negative_support)
    negative_f1 = 2 * negative_precision * negative_recall / max(1e-12, negative_precision + negative_recall)
    accuracy = (tp + tn) / len(valid)
    balanced_accuracy = (recall + negative_recall) / 2
    brier = statistics.fmean((score - actual) ** 2 for score, actual in zip(scores, actuals))
    base_rate = statistics.fmean(actuals)
    baseline_brier = statistics.fmean((base_rate - actual) ** 2 for actual in actuals)
    reliability = []
    ece = 0.0
    for bucket in range(10):
        low = bucket / 10
        high = (bucket + 1) / 10
        indexes = [
            index for index, score in enumerate(scores)
            if low <= score < high or (bucket == 9 and score == 1)
        ]
        if not indexes:
            continue
        predicted_rate = statistics.fmean(scores[index] for index in indexes)
        actual_rate = statistics.fmean(actuals[index] for index in indexes)
        ece += abs(predicted_rate - actual_rate) * len(indexes) / len(valid)
        reliability.append({
            "bucket": f"{bucket * 10}-{(bucket + 1) * 10}",
            "count": len(indexes),
            "predictedPct": round(predicted_rate * 100, 4),
            "actualPct": round(actual_rate * 100, 4),
        })
    slope, intercept = _platt_slope(actuals, scores)
    probability_std = statistics.pstdev(scores) if len(scores) > 1 else 0.0
    occupied_buckets = sum(1 for row in reliability if int(row.get("count") or 0) >= 30)
    probability_resolution_passed = probability_std >= 0.035 and occupied_buckets >= 4
    return {
        "available": True,
        "samples": len(valid),
        "independentDates": len({str(row.get("date")) for row in valid}),
        "positiveSupport": positive_support,
        "negativeSupport": negative_support,
        "accuracyPct": round(accuracy * 100, 5),
        "balancedAccuracyPct": round(balanced_accuracy * 100, 5),
        "precisionPct": round(precision * 100, 5),
        "recallPct": round(recall * 100, 5),
        "f1Pct": round(f1 * 100, 5),
        "macroF1Pct": round((f1 + negative_f1) / 2 * 100, 5),
        "rocAuc": None if (auc := _auc(actuals, scores)) is None else round(auc, 6),
        "prAuc": None if (pr_auc := _average_precision(actuals, scores)) is None else round(pr_auc, 6),
        "brier": round(brier, 6),
        "baselineBrier": round(baseline_brier, 6),
        "brierSkillScore": round(1 - brier / baseline_brier, 6) if baseline_brier > 1e-12 else None,
        "ecePct": round(ece * 100, 5),
        "calibrationSlope": None if slope is None else round(slope, 5),
        "calibrationIntercept": None if intercept is None else round(intercept, 5),
        "probabilityStd": round(probability_std, 6),
        "occupiedProbabilityBuckets": occupied_buckets,
        "probabilityResolutionPassed": probability_resolution_passed,
        "confusionMatrix": {"tp": tp, "tn": tn, "fp": fp, "fn": fn},
        "reliabilityCurve": reliability,
    }


def regression_metrics(rows: list[dict[str, Any]], prediction_key: str, actual_key: str = "actualReturn") -> dict[str, Any]:
    pairs = [
        (number(row.get(prediction_key)), number(row.get(actual_key)), row)
        for row in rows
        if row.get(prediction_key) is not None and row.get(actual_key) is not None
    ]
    if not pairs:
        return {"available": False, "reason": "No strict OOF regression rows.", "samples": 0}
    errors = [predicted - actual for predicted, actual, _ in pairs]
    actuals = [actual for _, actual, _ in pairs]
    mean_actual = statistics.fmean(actuals)
    ss_total = sum((actual - mean_actual) ** 2 for actual in actuals)
    ss_error = sum(error ** 2 for error in errors)
    direction = sum(1 for predicted, actual, _ in pairs if (predicted >= 0) == (actual >= 0)) / len(pairs)
    return {
        "available": True,
        "samples": len(pairs),
        "independentDates": len({str(row.get("date")) for _, _, row in pairs}),
        "mae": round(statistics.fmean(abs(error) for error in errors), 6),
        "rmse": round(math.sqrt(statistics.fmean(error ** 2 for error in errors)), 6),
        "r2": round(1 - ss_error / ss_total, 6) if ss_total > 1e-12 else None,
        "directionAccuracyPct": round(direction * 100, 5),
        "meanError": round(statistics.fmean(errors), 6),
        "errorP90": round(sorted(abs(error) for error in errors)[max(0, math.ceil(len(errors) * 0.9) - 1)], 6),
    }


def quantile_metrics(rows: list[dict[str, Any]]) -> dict[str, Any]:
    valid = [
        row for row in rows
        if all(row.get(key) is not None for key in ("quantileP10", "quantileP50", "quantileP90", "actualReturn"))
    ]
    if not valid:
        return {"available": False, "reason": "No strict OOF quantile rows.", "samples": 0}

    def pinball(alpha: float, predicted: float, actual: float) -> float:
        error = actual - predicted
        return max(alpha * error, (alpha - 1) * error)

    covered = [
        number(row["quantileP10"]) <= number(row["actualReturn"]) <= number(row["quantileP90"])
        for row in valid
    ]
    conformal = [
        row for row in valid
        if row.get("conformalP10") is not None and row.get("conformalP90") is not None
    ]
    return {
        "available": True,
        "samples": len(valid),
        "p10Pinball": round(statistics.fmean(pinball(0.1, number(row["quantileP10"]), number(row["actualReturn"])) for row in valid), 6),
        "p50Pinball": round(statistics.fmean(pinball(0.5, number(row["quantileP50"]), number(row["actualReturn"])) for row in valid), 6),
        "p90Pinball": round(statistics.fmean(pinball(0.9, number(row["quantileP90"]), number(row["actualReturn"])) for row in valid), 6),
        "intervalCoveragePct": round(statistics.fmean(covered) * 100, 5),
        "meanIntervalWidthPct": round(statistics.fmean(number(row["quantileP90"]) - number(row["quantileP10"]) for row in valid), 6),
        "conformalCoveragePct": (
            round(statistics.fmean(
                number(row["conformalP10"]) <= number(row["actualReturn"]) <= number(row["conformalP90"])
                for row in conformal
            ) * 100, 5)
            if conformal else None
        ),
    }


def rank_metrics(rows: list[dict[str, Any]], prediction_key: str = "rankerPrediction") -> dict[str, Any]:
    groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        if row.get(prediction_key) is not None and row.get("actualReturn") is not None:
            groups[str(row.get("date"))].append(row)
    correlations: list[float] = []
    top_returns: list[float] = []
    universe_returns: list[float] = []
    precision_hits: list[float] = []
    ndcg_values: list[float] = []
    top_target_rates: list[float] = []
    top_direction_rates: list[float] = []
    dated_top_returns: list[tuple[str, float]] = []
    for group_date, group in groups.items():
        if len(group) < 5:
            continue
        predicted_order = sorted(range(len(group)), key=lambda index: number(group[index].get(prediction_key)))
        actual_order = sorted(range(len(group)), key=lambda index: number(group[index].get("actualReturn")))
        predicted_rank = {row_index: rank for rank, row_index in enumerate(predicted_order)}
        actual_rank = {row_index: rank for rank, row_index in enumerate(actual_order)}
        center = (len(group) - 1) / 2
        numerator = sum((predicted_rank[index] - center) * (actual_rank[index] - center) for index in range(len(group)))
        denominator = sum((index - center) ** 2 for index in range(len(group)))
        correlations.append(numerator / denominator if denominator else 0.0)
        top_count = max(1, math.ceil(len(group) * 0.10))
        selected = set(predicted_order[-top_count:])
        actual_top = set(actual_order[-top_count:])
        precision_hits.append(len(selected & actual_top) / top_count)
        top_returns.append(statistics.fmean(number(group[index].get("actualReturn")) for index in selected))
        dated_top_returns.append((group_date, top_returns[-1]))
        universe_returns.append(statistics.fmean(number(row.get("actualReturn")) for row in group))
        top_target_rates.append(statistics.fmean(number(group[index].get("actualTarget")) >= 0.5 for index in selected))
        top_direction_rates.append(statistics.fmean(number(group[index].get("actualReturn")) > 0 for index in selected))
        relevance = [
            3 if number(row.get("actualTarget")) >= 0.5 and number(row.get("actualStop")) < 0.5
            else 0 if number(row.get("actualStop")) >= 0.5
            else 2 if number(row.get("actualReturn")) > 0
            else 1
            for row in group
        ]
        ranked_relevance = [relevance[index] for index in predicted_order[-top_count:][::-1]]
        ideal_relevance = sorted(relevance, reverse=True)[:top_count]
        dcg = sum((2 ** value - 1) / math.log2(position + 2) for position, value in enumerate(ranked_relevance))
        ideal_dcg = sum((2 ** value - 1) / math.log2(position + 2) for position, value in enumerate(ideal_relevance))
        ndcg_values.append(dcg / ideal_dcg if ideal_dcg > 0 else 0.0)
    if not correlations:
        return {"available": False, "reason": "Fewer than five symbols per OOF date.", "samples": 0}
    rank_ic = statistics.fmean(correlations)
    std = statistics.stdev(correlations) if len(correlations) > 1 else 0.0
    equity = peak = 1.0
    max_drawdown = 0.0
    for _, value in sorted(dated_top_returns):
        equity *= max(0.01, 1 + value / 100)
        peak = max(peak, equity)
        max_drawdown = max(max_drawdown, (peak - equity) / peak * 100)
    return {
        "available": True,
        "dateCount": len(correlations),
        "rankIc": round(rank_ic, 6),
        "rankIcNeweyWest95": newey_west_mean_ci(correlations),
        "icir": round(rank_ic / std, 6) if std > 1e-12 else None,
        "precisionAt10Pct": round(statistics.fmean(precision_hits) * 100, 5),
        "ndcgAt10Pct": round(statistics.fmean(ndcg_values), 6),
        "top10TargetFirstRatePct": round(statistics.fmean(top_target_rates) * 100, 5),
        "top10DirectionHitRatePct": round(statistics.fmean(top_direction_rates) * 100, 5),
        "topDecileNetReturnPct": round(statistics.fmean(top_returns), 6),
        "universeNetReturnPct": round(statistics.fmean(universe_returns), 6),
        "topDecileLiftPct": round(statistics.fmean(top_returns) - statistics.fmean(universe_returns), 6),
        "topDecileMaxDrawdownPct": round(max_drawdown, 6),
    }


def block_bootstrap_ci(
    rows: list[dict[str, Any]],
    statistic: Callable[[list[dict[str, Any]]], float | None],
    *,
    samples: int = 400,
    seed: int = 20260728,
) -> dict[str, Any]:
    groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        groups[str(row.get("date") or row.get("signalAt") or "unknown")].append(row)
    dates = sorted(groups)
    if len(dates) < 5:
        return {"available": False, "reason": "Fewer than five independent date blocks.", "dateBlocks": len(dates)}
    rng = random.Random(seed)
    values = []
    for _ in range(samples):
        sampled = []
        for _ in dates:
            sampled.extend(groups[rng.choice(dates)])
        value = statistic(sampled)
        if value is not None and math.isfinite(value):
            values.append(value)
    if len(values) < max(20, samples // 4):
        return {"available": False, "reason": "Bootstrap statistic was not identifiable.", "dateBlocks": len(dates)}
    values.sort()
    return {
        "available": True,
        "dateBlocks": len(dates),
        "samples": len(values),
        "low": round(values[int(len(values) * 0.025)], 6),
        "high": round(values[min(len(values) - 1, int(len(values) * 0.975))], 6),
    }


def classification_accuracy_block_ci(
    rows: list[dict[str, Any]],
    probability_key: str,
    *,
    actual_key: str = "actualTarget",
    samples: int = 400,
    seed: int = 20260728,
) -> dict[str, Any]:
    groups: dict[str, tuple[int, int]] = {}
    for row in rows:
        probability = row.get(probability_key)
        if probability is None:
            continue
        day = str(row.get("date") or row.get("signalAt") or "unknown")
        correct, total = groups.get(day, (0, 0))
        predicted = 1 if number(probability) >= 0.5 else 0
        actual = 1 if number(row.get(actual_key)) >= 0.5 else 0
        groups[day] = (correct + int(predicted == actual), total + 1)
    dates = sorted(groups)
    if len(dates) < 5:
        return {"available": False, "reason": "Fewer than five independent date blocks.", "dateBlocks": len(dates)}
    rng = random.Random(seed)
    values = []
    for _ in range(samples):
        correct = 0
        total = 0
        for _ in dates:
            day_correct, day_total = groups[rng.choice(dates)]
            correct += day_correct
            total += day_total
        if total:
            values.append(correct / total * 100.0)
    values.sort()
    return {
        "available": bool(values),
        "dateBlocks": len(dates),
        "samples": len(values),
        "low": round(values[int(len(values) * 0.025)], 6) if values else None,
        "high": round(values[min(len(values) - 1, int(len(values) * 0.975))], 6) if values else None,
    }


def read_oof_rows(path: Path, manifest: dict[str, Any], market: str) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    malformed = 0
    try:
        with gzip.open(path, "rt", encoding="utf-8") as handle:
            for line in handle:
                try:
                    row = json.loads(line)
                    row["predictionId"] = prediction_id(row, manifest)
                    rows.append(row)
                except (ValueError, TypeError):
                    malformed += 1
    except OSError:
        return [], {"available": False, "reason": "OOF artifact is unreadable.", "path": str(path)}
    counts = Counter(row["predictionId"] for row in rows)
    duplicates = sum(count - 1 for count in counts.values() if count > 1)
    cross_market = sum(1 for row in rows if str(row.get("market") or "").upper() != market)
    future_rows = sum(
        1 for row in rows
        if row.get("availableAt") and row.get("signalAt") and str(row["availableAt"]) > str(row["signalAt"])
    )
    clean = []
    seen = set()
    for row in rows:
        if row["predictionId"] in seen or str(row.get("market") or "").upper() != market:
            continue
        seen.add(row["predictionId"])
        clean.append(row)
    return clean, {
        "available": True,
        "rawRows": len(rows),
        "uniqueRows": len(clean),
        "duplicateRows": duplicates,
        "crossMarketRows": cross_market,
        "futureAvailabilityRows": future_rows,
        "malformedRows": malformed,
        "independentDates": len({str(row.get("date")) for row in clean}),
        "symbols": len({str(row.get("symbol")) for row in clean}),
        "path": str(path),
    }


def _model_status(model: dict[str, Any]) -> str:
    if model.get("productionEvidencePassed"):
        return "shadow_candidate"
    if model.get("available"):
        return "research"
    return "evidence_insufficient"


def _hard_gate(dataset: dict[str, Any], horizon_model: dict[str, Any], audit: dict[str, Any]) -> dict[str, Any]:
    metrics = horizon_model.get("metrics") or {}
    direction_metrics = horizon_model.get("directionMetrics") or {}
    long_gate = horizon_model.get("longTradeGate") or {}
    long_evidence = long_gate.get("testEvidence") or {}
    folds = horizon_model.get("foldMetrics") or []
    checks = [
        ("rows", number(horizon_model.get("rowCount")) >= PRODUCTION_THRESHOLDS["minRowsPerHorizon"], number(horizon_model.get("rowCount")), PRODUCTION_THRESHOLDS["minRowsPerHorizon"]),
        ("oof_rows", number(horizon_model.get("metaTestRows")) >= PRODUCTION_THRESHOLDS["minOofTestRows"], number(horizon_model.get("metaTestRows")), PRODUCTION_THRESHOLDS["minOofTestRows"]),
        ("test_dates", number(metrics.get("testDates", audit.get("independentDates"))) >= PRODUCTION_THRESHOLDS["minIndependentTestDates"], number(metrics.get("testDates", audit.get("independentDates"))), PRODUCTION_THRESHOLDS["minIndependentTestDates"]),
        ("target_events", number(horizon_model.get("eventCounts", {}).get("target")) >= PRODUCTION_THRESHOLDS["minTargetEvents"], number(horizon_model.get("eventCounts", {}).get("target")), PRODUCTION_THRESHOLDS["minTargetEvents"]),
        ("stop_events", number(horizon_model.get("eventCounts", {}).get("stop")) >= PRODUCTION_THRESHOLDS["minStopEvents"], number(horizon_model.get("eventCounts", {}).get("stop")), PRODUCTION_THRESHOLDS["minStopEvents"]),
        ("folds", len(folds) >= PRODUCTION_THRESHOLDS["minFolds"], len(folds), PRODUCTION_THRESHOLDS["minFolds"]),
        ("positive_folds", sum(1 for fold in folds if fold.get("positive")) >= PRODUCTION_THRESHOLDS["minPositiveFolds"], sum(1 for fold in folds if fold.get("positive")), PRODUCTION_THRESHOLDS["minPositiveFolds"]),
        ("duplicates", int(audit.get("duplicateRows") or 0) == 0, int(audit.get("duplicateRows") or 0), 0),
        ("cross_market", int(audit.get("crossMarketRows") or 0) == 0, int(audit.get("crossMarketRows") or 0), 0),
        ("future_leakage", int(audit.get("futureAvailabilityRows") or 0) == 0 and int(dataset.get("pointInTimeJoinViolationCount") or 0) == 0, int(audit.get("futureAvailabilityRows") or 0) + int(dataset.get("pointInTimeJoinViolationCount") or 0), 0),
        ("brier_skill", number(metrics.get("brierSkillScore"), -1) > PRODUCTION_THRESHOLDS["minBrierSkill"], number(metrics.get("brierSkillScore"), -1), f">{PRODUCTION_THRESHOLDS['minBrierSkill']}"),
        ("direction_brier_skill", number(direction_metrics.get("brierSkillScore"), -1) > PRODUCTION_THRESHOLDS["minBrierSkill"], number(direction_metrics.get("brierSkillScore"), -1), f">{PRODUCTION_THRESHOLDS['minBrierSkill']}"),
        ("direction_balanced_accuracy", number(direction_metrics.get("balancedAccuracyPct")) > 50.0, number(direction_metrics.get("balancedAccuracyPct")), ">50%"),
        ("eligible_long_accuracy", long_gate.get("active") is True and number(long_evidence.get("directionHitRatePct")) >= 57.0 and number(long_evidence.get("directionHitRate95LowerPct")) > 50.0, {"hitRatePct": long_evidence.get("directionHitRatePct"), "ci95LowerPct": long_evidence.get("directionHitRate95LowerPct")}, "hit>=57% and block-CI lower>50%"),
        ("eligible_long_support", number(long_evidence.get("signalCount")) >= 200 and number(long_evidence.get("signalDates")) >= 80, {"signals": long_evidence.get("signalCount"), "dates": long_evidence.get("signalDates")}, "200 signals / 80 dates"),
        ("ece", number(metrics.get("ecePct"), 100) <= PRODUCTION_THRESHOLDS["maxEcePct"], number(metrics.get("ecePct"), 100), PRODUCTION_THRESHOLDS["maxEcePct"]),
        ("calibration_slope", PRODUCTION_THRESHOLDS["minCalibrationSlope"] <= number(metrics.get("calibrationSlope"), -1) <= PRODUCTION_THRESHOLDS["maxCalibrationSlope"], number(metrics.get("calibrationSlope"), -1), "0.8-1.2"),
        ("probability_bucket", number(metrics.get("probabilityBucketMinCount")) >= PRODUCTION_THRESHOLDS["minProbabilityBucketEvents"], number(metrics.get("probabilityBucketMinCount")), PRODUCTION_THRESHOLDS["minProbabilityBucketEvents"]),
        ("target_probability_resolution", metrics.get("probabilityResolutionPassed") is True, {"std": metrics.get("probabilityStd"), "buckets": metrics.get("occupiedProbabilityBuckets")}, "std>=0.035 and >=4 buckets"),
        ("direction_probability_resolution", direction_metrics.get("probabilityResolutionPassed") is True, {"std": direction_metrics.get("probabilityStd"), "buckets": direction_metrics.get("occupiedProbabilityBuckets")}, "std>=0.035 and >=4 buckets"),
        ("top_k_lift", number(horizon_model.get("rankingMetrics", {}).get("topDecileLift")) > 0, number(horizon_model.get("rankingMetrics", {}).get("topDecileLift")), ">0"),
        ("top_k_absolute_return", number(horizon_model.get("rankingMetrics", {}).get("topDecileNetReturn")) > 0, number(horizon_model.get("rankingMetrics", {}).get("topDecileNetReturn")), ">0"),
        ("net_ev", number(horizon_model.get("longTradeExpectedValue", {}).get("expectedValuePct")) > 0, number(horizon_model.get("longTradeExpectedValue", {}).get("expectedValuePct")), ">0 on eligible long signals"),
        ("feature_drift", max([number(fold.get("featureDrift", {}).get("maxPsi")) for fold in folds] or [0]) <= PRODUCTION_THRESHOLDS["maxFeaturePsi"], max([number(fold.get("featureDrift", {}).get("maxPsi")) for fold in folds] or [0]), PRODUCTION_THRESHOLDS["maxFeaturePsi"]),
    ]
    rows = [{"id": key, "passed": passed, "value": value, "threshold": threshold} for key, passed, value, threshold in checks]
    return {
        "passed": bool(rows) and all(row["passed"] for row in rows),
        "checks": rows,
        "failedChecks": [row["id"] for row in rows if not row["passed"]],
    }


def horizon_report(
    market: str,
    registry: dict[str, Any],
    model: dict[str, Any],
    oof_dir: Path,
) -> dict[str, Any]:
    manifest = registry.get("manifest") or {}
    artifact = model.get("oofArtifact") or {}
    artifact_path = oof_dir / str(artifact.get("filename") or "")
    rows: list[dict[str, Any]] = []
    audit = {"available": False, "reason": model.get("reason") or "No usable OOF artifact."}
    if artifact.get("filename") and artifact_path.exists():
        rows, audit = read_oof_rows(artifact_path, manifest, market)
    classifiers = []
    for key, label, actual_key in MODEL_OUTPUTS:
        result = classification_metrics(rows, key, actual_key)
        if result.get("available"):
            result["accuracyCi95"] = classification_accuracy_block_ci(rows, key, actual_key=actual_key)
        classifiers.append({"id": key, "name": label, "task": "classification", "target": actual_key, "metrics": result})
    quantiles = quantile_metrics(rows)
    regression = regression_metrics(rows, "quantileP50")
    ranking_key = "selectionScore" if any(row.get("selectionScore") is not None for row in rows) else "rankerPrediction"
    ranking = rank_metrics(rows, ranking_key)
    ranking["scoreKey"] = ranking_key
    hard_gate = _hard_gate(registry.get("dataset") or {}, model, audit)
    return {
        "modelId": model.get("modelVersion") or f"{manifest.get('model_version', market.lower())}-{model.get('horizon', 'unknown')}d",
        "modelVersion": manifest.get("model_version"),
        "market": market,
        "horizon": model.get("horizon"),
        "family": "market_multitask",
        "status": _model_status(model),
        "available": bool(model.get("available") and rows),
        "reason": model.get("reason") or ("" if rows else "OOF evidence was not usable."),
        "sampleAudit": audit,
        "classifiers": classifiers,
        "ranking": ranking,
        "regression": regression,
        "quantiles": quantiles,
        "weights": model.get("weights") or [],
        "prunedModels": model.get("prunedModels") or [],
        "foldMetrics": model.get("foldMetrics") or model.get("folds") or [],
        "featureImportance": model.get("featureImportance") or [],
        "featureAblation": model.get("featureAblation") or {},
        "diagnosticBuckets": model.get("diagnosticBuckets") or {},
        "modelComparison": model.get("modelComparison") or [],
        "selectiveRankingHead": model.get("selectiveRankingHead") or {},
        "longTradeGate": model.get("longTradeGate") or {},
        "calibrator": model.get("calibrator"),
        "expectedValue": model.get("expectedValue"),
        "longTradeExpectedValue": model.get("longTradeExpectedValue"),
        "hardGate": hard_gate,
    }


def factor_model_reports(root: Path, market: str) -> list[dict[str, Any]]:
    rows = []
    for path in sorted((root / ".cache" / "models" / "factor-research").glob("*.json")):
        payload = read_json(path, {})
        if str(payload.get("market") or "").upper() != market:
            continue
        scope = str(payload.get("scope") or ("market" if path.stem.startswith("cross-sectional-") else "stock"))
        if scope == "market":
            horizon_rows = payload.get("horizonResults") or []
            tests = [row.get("mlBacktest", {}).get("test", {}) for row in horizon_rows if isinstance(row, dict)]
            tests = [row for row in tests if row]
            sample_count = sum(int(number(row.get("rowCount"))) for row in horizon_rows)
            local_gate_passed = bool(payload.get("eligibleForLiveWeight") is True and any(row.get("mlBacktest", {}).get("active") is True for row in horizon_rows))
            rows.append({
                "modelId": path.stem,
                "modelVersion": payload.get("evidenceId") or payload.get("savedAt"),
                "market": market,
                "scope": "market",
                "horizon": payload.get("horizons"),
                "family": "factor_research",
                "status": "shadow" if local_gate_passed else "research",
                "available": bool(horizon_rows),
                "sampleCount": sample_count,
                "symbolCount": payload.get("symbolCount"),
                "eligibleForLiveWeight": False,
                "localOosGatePassed": local_gate_passed,
                "metrics": {
                    "rankIc": metric(sum(number(row.get("rank_ic")) for row in tests) / len(tests) if tests else None, reason="No market-level untouched Rank IC."),
                    "directionAccuracyPct": metric(sum(number(row.get("direction_hit_rate_pct")) for row in tests) / len(tests) if tests else None, reason="No market-level untouched direction result."),
                    "netSignalReturnPct": metric(sum(number(row.get("avg_net_signal_return_pct")) for row in tests) / len(tests) if tests else None, reason="No market-level cost-adjusted factor result."),
                },
                "hardGate": {"passed": False, "failedChecks": [*([] if local_gate_passed else ["market_factor_oos_gate_failed"]), "prediction_champion_required"]},
            })
            continue
        backtest = payload.get("mlBacktest") or {}
        policy_version = int(number(payload.get("admissionPolicyVersion")))
        local_gate_passed = bool(
            policy_version >= 2
            and payload.get("eligibleForLiveWeight") is True
            and backtest.get("active") is True
        )
        # A per-symbol factor experiment is research evidence, not a production
        # market model. It remains zero-weight until strict market-level OOF is
        # independently eligible.
        live_eligible = False
        legacy_quarantined = policy_version < 2
        failed_checks = ["market_level_oof_required"]
        if legacy_quarantined:
            failed_checks.insert(0, "legacy_admission_policy_quarantined")
        elif not local_gate_passed:
            failed_checks.insert(0, "strict_factor_oos_gate_failed")
        rows.append({
            "modelId": path.stem,
            "modelVersion": payload.get("savedAt"),
            "market": market,
            "symbol": payload.get("symbol"),
            "scope": "stock",
            "horizon": payload.get("horizonDays"),
            "family": "factor_research",
            "status": "quarantined" if legacy_quarantined else ("research" if payload.get("sampleCount") else "evidence_insufficient"),
            "available": bool(payload.get("sampleCount")),
            "sampleCount": payload.get("sampleCount"),
            "candidateCount": payload.get("candidateCount"),
            "admissionPolicyVersion": policy_version,
            "legacyQuarantined": legacy_quarantined,
            "eligibleForLiveWeight": live_eligible,
            "localOosGatePassed": local_gate_passed,
            "reportedAdmittedCount": payload.get("admittedCount"),
            "admittedCount": payload.get("admittedCount") if live_eligible else 0,
            "metrics": {
                "ic": metric(backtest.get("ic"), reason="Factor IC was not persisted."),
                "rankIc": metric(backtest.get("rank_ic", backtest.get("rankIc")), reason="Factor Rank IC was not persisted."),
                "directionAccuracyPct": metric(backtest.get("direction_hit_rate_pct", backtest.get("directionHitRate")), reason="No untouched holdout result."),
            },
            "hardGate": {"passed": False, "failedChecks": failed_checks},
        })
    return rows


def intraday_model_report(root: Path, market: str) -> dict[str, Any]:
    payload = read_json(root / ".cache" / "backend-monitor" / f"intraday-model-{market.lower()}.json", {})
    holdout = payload.get("holdout") or payload.get("test") or {}
    samples = int(number(payload.get("sampleCount", payload.get("samples"))))
    return {
        "modelId": f"{market.lower()}-intraday",
        "modelVersion": payload.get("updatedAt") or payload.get("trainedAt"),
        "market": market,
        "family": "intraday",
        "status": "research" if samples else "evidence_insufficient",
        "available": bool(samples),
        "sampleCount": samples,
        "metrics": {
            "accuracyPct": metric(holdout.get("directionalAccuracy", holdout.get("direction_hit_rate_pct")), reason="No untouched intraday holdout."),
            "mae": metric(holdout.get("mae"), reason="MAE was not persisted."),
            "rmse": metric(holdout.get("rmse"), reason="RMSE was not persisted."),
            "rankIc": metric(holdout.get("rank_ic", holdout.get("rankIc")), reason="Rank IC was not persisted."),
        },
        "hardGate": {"passed": False, "failedChecks": ["intraday_production_gate_not_implemented"]},
    }


def paper_agent_reports(root: Path, market: str) -> list[dict[str, Any]]:
    try:
        state = load_paper_agent_state(market, str(root / ".cache" / "quant-control-plane.sqlite3"))
    except Exception as exc:
        return [{
            "modelId": f"{market.lower()}-paper-agent-store",
            "market": market,
            "family": "paper_agent",
            "status": "unavailable",
            "available": False,
            "reason": str(exc)[:240],
        }]
    reports = []
    for agent in state.get("ledger", {}).get("agents", []):
        trades = list(agent.get("trades") or [])
        closed = [row for row in trades if str(row.get("side") or "").lower() == "sell"]
        returns = [number(row.get("returnPct", row.get("reward"))) for row in closed if row.get("returnPct", row.get("reward")) is not None]
        wins = sum(1 for value in returns if value > 0)
        losses = [abs(value) for value in returns if value < 0]
        gains = [value for value in returns if value > 0]
        reports.append({
            "modelId": f"{market.lower()}-paper-{agent.get('id')}",
            "modelVersion": state.get("revision"),
            "market": market,
            "family": "paper_agent",
            "status": "paper",
            "available": True,
            "name": agent.get("name"),
            "style": agent.get("style"),
            "metrics": {
                "netReturnPct": round(number(agent.get("returnPct")), 6),
                "tradeCount": int(number(agent.get("stats", {}).get("trades"), len(trades))),
                "closedTrades": len(closed),
                "winRatePct": round(wins / len(returns) * 100, 5) if returns else None,
                "profitFactor": round(sum(gains) / sum(losses), 6) if losses else None,
                "sharpe": None,
                "sortino": None,
                "maxDrawdownPct": None,
                "turnoverPct": None,
                "capacity": "not_established",
            },
            "hardGate": {"passed": False, "failedChecks": ["paper_observation_and_risk_metrics_incomplete"]},
        })
    return reports


def supervisor_reports(root: Path, market: str) -> list[dict[str, Any]]:
    state = read_json(root / ".cache" / "training-supervisor" / "state.json", {})
    market_state = (state.get("markets") or {}).get(market) or {}
    reports = []
    for provider in ("openai", "siliconflow", "hunyuan"):
        review = next((row for row in market_state.get("reviewers") or [] if row.get("provider") == provider), {})
        reports.append({
            "modelId": f"{market.lower()}-supervisor-{provider}",
            "market": market,
            "family": "ai_supervisor",
            "status": "available" if review.get("available") else "unavailable",
            "available": bool(review.get("available")),
            "provider": provider,
            "model": review.get("model"),
            "verdict": review.get("verdict") or "not_reviewed",
            "score": review.get("score"),
            "rationale": review.get("rationale") or review.get("error") or "No review evidence.",
            "blockingIssues": review.get("blockingIssues") or [],
            "recommendedActions": review.get("recommendedActions") or [],
            "reviewedAt": review.get("reviewedAt"),
            "note": "Supervisor opinions are governance evidence, not prediction accuracy.",
        })
    return reports


def latest_registry(root: Path, market: str) -> dict[str, Any] | None:
    directory = root / ".cache" / "models" / "registry" / market.lower()
    index = read_json(directory / "index.json", {})
    filename = (index.get("latest") or {}).get("filename")
    if not filename:
        return None
    payload = read_json(directory / filename, None)
    return payload if isinstance(payload, dict) else None


def market_report(root: Path, market: str) -> dict[str, Any]:
    registry = latest_registry(root, market)
    horizon_reports: list[dict[str, Any]] = []
    reviewers = supervisor_reports(root, market)
    reviewer_approvals = sum(
        1 for row in reviewers
        if row.get("available") and row.get("verdict") == "accept"
    )
    if registry:
        for model in registry.get("horizonModels") or []:
            report = horizon_report(
                market,
                registry,
                model,
                root / ".cache" / "models" / "oof" / market.lower(),
            )
            review_check = {
                "id": "ai_supervisor_consensus",
                "passed": reviewer_approvals >= 2,
                "value": reviewer_approvals,
                "threshold": 2,
            }
            report["hardGate"]["checks"].append(review_check)
            if not review_check["passed"]:
                report["hardGate"]["passed"] = False
                report["hardGate"]["failedChecks"].append(review_check["id"])
            horizon_reports.append(report)
    dataset = registry.get("dataset") if registry else {}
    production = registry.get("productionEligibility") if registry else {}
    models = [
        *horizon_reports,
        intraday_model_report(root, market),
        *factor_model_reports(root, market),
        *paper_agent_reports(root, market),
        *reviewers,
    ]
    return {
        "market": market,
        "registryAvailable": registry is not None,
        "modelVersion": (registry or {}).get("manifest", {}).get("model_version"),
        "deploymentStatus": (registry or {}).get("manifest", {}).get("deployment_status", "research"),
        "dataset": dataset or {},
        "trainedHorizons": (registry or {}).get("trainedHorizons") or [],
        "withheldHorizons": (registry or {}).get("withheldHorizons") or [],
        "limitedDataPolicy": (registry or {}).get("limitedDataPolicy") or {},
        "productionEligibility": production or {
            "eligible": False,
            "reason": "No market-level immutable model registry exists.",
        },
        "models": models,
        "counts": {
            "total": len(models),
            "available": sum(1 for model in models if model.get("available")),
            "hardGatePassed": sum(1 for model in models if model.get("hardGate", {}).get("passed")),
            "aiSupervisorApprovals": reviewer_approvals,
        },
    }


def _job_header(path: Path) -> dict[str, Any]:
    """Read bounded top-level job metadata without loading multi-megabyte results."""
    try:
        with path.open("rb") as handle:
            text = handle.read(8_192).decode("utf-8", errors="ignore")
    except OSError:
        return {}

    def string_value(key: str) -> str | None:
        match = re.search(rf'"{re.escape(key)}"\s*:\s*"([^"\\]*(?:\\.[^"\\]*)*)"', text)
        if not match:
            return None
        try:
            return json.loads(f'"{match.group(1)}"')
        except (ValueError, TypeError):
            return match.group(1)

    def integer_value(key: str) -> int:
        match = re.search(rf'"{re.escape(key)}"\s*:\s*(-?\d+)', text)
        return int(match.group(1)) if match else 0

    return {
        "id": string_value("id") or path.stem,
        "type": string_value("type") or "unknown",
        "runtimeVersion": integer_value("runtimeVersion"),
        "market": string_value("market"),
        "status": string_value("status") or "unknown",
        "failureCategory": string_value("failureCategory"),
        "createdAt": string_value("createdAt"),
        "updatedAt": string_value("updatedAt"),
    }


def _job_summary(root: Path) -> dict[str, Any]:
    current_runtime_version = 3
    job_paths = sorted(
        (root / ".cache" / "background-jobs").glob("backtest-*.json"),
        key=lambda path: path.name,
        reverse=True,
    )
    audit_limit = max(50, min(1_000, int(os.getenv("MODEL_REPORT_JOB_AUDIT_LIMIT", "400"))))
    audited_paths = job_paths[:audit_limit]
    counts = Counter()
    types = Counter()
    failure_categories = Counter()
    current_counts = Counter()
    current_failures = Counter()
    by_runtime: dict[str, Counter] = {}
    recent_rows: list[dict[str, Any]] = []
    for path in audited_paths:
        payload = _job_header(path)
        if not payload:
            continue
        status = str(payload.get("status") or "unknown")
        counts[status] += 1
        types[str(payload.get("type") or "unknown")] += 1
        recent_rows.append(payload)
        runtime_version = int(number(payload.get("runtimeVersion")))
        runtime_key = str(runtime_version or "legacy")
        by_runtime.setdefault(runtime_key, Counter())[status] += 1
        if runtime_version == current_runtime_version:
            current_counts[status] += 1
        if status == "failed":
            failure_categories[str(payload.get("failureCategory") or "unclassified")] += 1
            if runtime_version == current_runtime_version:
                current_failures[str(payload.get("failureCategory") or "unclassified")] += 1
    terminal = counts["complete"] + counts["failed"]
    current_terminal = current_counts["complete"] + current_counts["failed"]
    recent_rows.sort(key=lambda row: str(row.get("updatedAt") or row.get("createdAt") or ""), reverse=True)
    recent_terminal_rows = [row for row in recent_rows if row.get("status") in {"complete", "failed"}][:50]
    recent_failures = sum(1 for row in recent_terminal_rows if row.get("status") == "failed")
    return {
        "total": sum(counts.values()),
        "totalArtifacts": len(job_paths),
        "auditedJobs": len(audited_paths),
        "auditCoveragePct": round(len(audited_paths) / len(job_paths) * 100, 5) if job_paths else 100.0,
        "status": dict(counts),
        "types": dict(types),
        "failureCategories": dict(failure_categories),
        "terminalFailureRatePct": round(counts["failed"] / terminal * 100, 5) if terminal else None,
        "currentRuntime": {
            "version": current_runtime_version,
            "status": dict(current_counts),
            "failureCategories": dict(current_failures),
            "terminalFailureRatePct": round(current_counts["failed"] / current_terminal * 100, 5) if current_terminal else None,
        },
        "byRuntime": {key: dict(value) for key, value in sorted(by_runtime.items())},
        "recent50": {
            "terminalJobs": len(recent_terminal_rows),
            "failedJobs": recent_failures,
            "failureRatePct": round(recent_failures / len(recent_terminal_rows) * 100, 5) if recent_terminal_rows else None,
        },
        "interpretation": "Training reliability uses the newest bounded backtest audit window so report generation never reloads every historical result blob. All artifacts remain retained; currentRuntime isolates Runtime V3 fold-checkpoint jobs.",
    }


def build_report_evidence(root: Path, markets: list[str] | None = None) -> dict[str, Any]:
    selected = [market for market in (markets or list(MARKETS)) if market in MARKETS]
    generated_at = iso_now()
    market_rows = [market_report(root, market) for market in selected]
    report_id = f"model-training-report-{generated_at[:19].replace(':', '').replace('-', '')}-{stable_hash([row['modelVersion'] for row in market_rows], 8)}"
    blockers = []
    for row in market_rows:
        if not row["registryAvailable"]:
            blockers.append(f"{row['market']}: no market-level registry")
        for model in row["models"]:
            if model.get("family") == "market_multitask" and not model.get("hardGate", {}).get("passed"):
                blockers.append(f"{row['market']} {model.get('horizon')}d: {', '.join(model.get('hardGate', {}).get('failedChecks') or ['insufficient evidence'])}")
    return {
        "schemaVersion": 2,
        "reportId": report_id,
        "generatedAt": generated_at,
        "title": "Global Quant Watch 全模型训练与生产验收报告",
        "scope": selected,
        "honestBoundary": "Unavailable metrics remain unavailable. Duplicate decisions, paper returns, rule heads, and AI reviewer opinions are not relabeled as model accuracy.",
        "productionReady": all(
            row.get("productionEligibility", {}).get("eligible") is True
            and any(model.get("hardGate", {}).get("passed") for model in row["models"])
            for row in market_rows
        ),
        "thresholds": PRODUCTION_THRESHOLDS,
        "jobReliability": _job_summary(root),
        "markets": market_rows,
        "blockers": blockers,
        "recommendedActions": [
            "Quarantine cross-market and duplicate prediction rows before any metric is recomputed.",
            "Complete point-in-time universe, delisting, corporate-action, and event histories.",
            "Run market-level purged walk-forward OOF training to the production sample thresholds.",
            "Keep every failed or under-supported model in Research/Shadow with zero live ensemble weight.",
            "Require deterministic gates before requesting AI supervisor review.",
        ],
    }


def _fmt(value: Any, digits: int = 2, suffix: str = "") -> str:
    if value is None:
        return "证据不足"
    try:
        return f"{float(value):.{digits}f}{suffix}"
    except (TypeError, ValueError):
        return str(value)


def _reliability_svg(metrics: dict[str, Any]) -> str:
    points = metrics.get("reliabilityCurve") or []
    if not points:
        return ""
    width, height, pad = 420, 210, 32
    coords = []
    for row in points:
        x = pad + clamp(row.get("predictedPct"), 0, 100) / 100 * (width - pad * 2)
        y = height - pad - clamp(row.get("actualPct"), 0, 100) / 100 * (height - pad * 2)
        coords.append(f"{x:.1f},{y:.1f}")
    return f"""
      <figure class="evidence-chart">
        <figcaption>OOF 概率可靠性曲线</figcaption>
        <svg viewBox="0 0 {width} {height}" role="img" aria-label="OOF reliability curve">
          <line x1="{pad}" y1="{height-pad}" x2="{width-pad}" y2="{pad}" class="ideal"/>
          <polyline points="{' '.join(coords)}" class="actual"/>
          {''.join(f'<circle cx="{point.split(",")[0]}" cy="{point.split(",")[1]}" r="4"/>' for point in coords)}
          <text x="{pad}" y="{height-9}">0%</text><text x="{width-pad-25}" y="{height-9}">100%</text>
        </svg>
      </figure>
    """


def _confusion_html(metrics: dict[str, Any]) -> str:
    values = metrics.get("confusionMatrix") or {}
    if not values:
        return ""
    return f"""
      <figure class="evidence-chart confusion">
        <figcaption>OOF 混淆矩阵</figcaption>
        <div class="confusion-grid">
          <span><small>TN</small><b>{int(number(values.get('tn')))}</b></span>
          <span><small>FP</small><b>{int(number(values.get('fp')))}</b></span>
          <span><small>FN</small><b>{int(number(values.get('fn')))}</b></span>
          <span><small>TP</small><b>{int(number(values.get('tp')))}</b></span>
        </div>
      </figure>
    """


def render_html(evidence: dict[str, Any], target: Path) -> None:
    market_sections = []
    for market in evidence["markets"]:
        dataset = market.get("dataset") or {}
        cards = []
        for model in market["models"]:
            family = model.get("family")
            title = model.get("name") or model.get("modelId")
            if family == "market_multitask":
                primary = next((row for row in model.get("classifiers") or [] if row["id"] == "directionProbability"), None) or next((row for row in model.get("classifiers") or [] if row["id"] == "ensembleProbability"), None)
                metrics = (primary or {}).get("metrics") or {}
                metric_rows = [
                    ("Accuracy", _fmt(metrics.get("accuracyPct"), suffix="%")),
                    ("Precision", _fmt(metrics.get("precisionPct"), suffix="%")),
                    ("Recall", _fmt(metrics.get("recallPct"), suffix="%")),
                    ("F1", _fmt(metrics.get("f1Pct"), suffix="%")),
                    ("ROC-AUC", _fmt(metrics.get("rocAuc"), 3)),
                    ("Brier Skill", _fmt(metrics.get("brierSkillScore"), 3)),
                    ("ECE", _fmt(metrics.get("ecePct"), suffix="%")),
                    ("概率分辨率", "通过" if metrics.get("probabilityResolutionPassed") else f"不足 · σ={_fmt(metrics.get('probabilityStd'), 3)} / {int(number(metrics.get('occupiedProbabilityBuckets')))}桶"),
                    ("独立日期", _fmt(model.get("sampleAudit", {}).get("independentDates"), 0)),
                ]
                visuals = _reliability_svg(metrics) + _confusion_html(metrics)
                task_rows = []
                for classifier in model.get("classifiers") or []:
                    values = classifier.get("metrics") or {}
                    task_rows.append(
                        f"<tr><td>{html.escape(str(classifier.get('name')))}</td><td>分类</td>"
                        f"<td>{_fmt(values.get('samples'), 0)}</td><td>{_fmt(values.get('accuracyPct'), suffix='%')}</td>"
                        f"<td>{_fmt(values.get('f1Pct'), suffix='%')}</td><td>{_fmt(values.get('brierSkillScore'), 3)}</td></tr>"
                    )
                ranking = model.get("ranking") or {}
                regression = model.get("regression") or {}
                quantiles = model.get("quantiles") or {}
                task_rows.extend([
                    f"<tr><td>横截面排序</td><td>排序</td><td>{_fmt(ranking.get('dateCount'), 0)}</td><td>IC {_fmt(ranking.get('rankIc'), 3)} / NDCG {_fmt(ranking.get('ndcgAt10Pct'), 3)}</td><td>目标先到 {_fmt(ranking.get('top10TargetFirstRatePct'), suffix='%')} / 方向 {_fmt(ranking.get('top10DirectionHitRatePct'), suffix='%')}</td><td>Lift {_fmt(ranking.get('topDecileLiftPct'), 3)} / DD {_fmt(ranking.get('topDecileMaxDrawdownPct'), suffix='%')}</td></tr>",
                    f"<tr><td>收益中位数</td><td>回归</td><td>{_fmt(regression.get('samples'), 0)}</td><td>{_fmt(regression.get('directionAccuracyPct'), suffix='%')}</td><td>{_fmt(regression.get('mae'), 3)}</td><td>{_fmt(regression.get('r2'), 3)}</td></tr>",
                    f"<tr><td>收益分位数</td><td>Quantile</td><td>{_fmt(quantiles.get('samples'), 0)}</td><td>{_fmt(quantiles.get('intervalCoveragePct'), suffix='%')}</td><td>{_fmt(quantiles.get('p50Pinball'), 3)}</td><td>{_fmt(quantiles.get('meanIntervalWidthPct'), 3)}</td></tr>",
                ])
                task_detail = (
                    "<details class=\"task-evidence\"><summary>查看全部任务指标</summary>"
                    "<table><thead><tr><th>模型/输出</th><th>任务</th><th>样本</th><th>主指标</th><th>次指标</th><th>校准/增益</th></tr></thead>"
                    f"<tbody>{''.join(task_rows)}</tbody></table></details>"
                )
            elif family == "paper_agent":
                values = model.get("metrics") or {}
                metric_rows = [
                    ("成本后收益", _fmt(values.get("netReturnPct"), suffix="%")),
                    ("成交", _fmt(values.get("tradeCount"), 0)),
                    ("胜率", _fmt(values.get("winRatePct"), suffix="%")),
                    ("Profit Factor", _fmt(values.get("profitFactor"), 2)),
                ]
                visuals = ""
                task_detail = ""
            elif family == "ai_supervisor":
                metric_rows = [("结论", model.get("verdict")), ("评分", _fmt(model.get("score"), 0)), ("可用", "是" if model.get("available") else "否")]
                visuals = ""
                task_detail = ""
            else:
                values = model.get("metrics") or {}
                metric_rows = [(key, _fmt(value.get("value") if isinstance(value, dict) else value)) for key, value in list(values.items())[:8]]
                visuals = ""
                task_detail = ""
            rows_html = "".join(f"<div><span>{html.escape(str(label))}</span><strong>{html.escape(str(value))}</strong></div>" for label, value in metric_rows)
            gate = model.get("hardGate") or {}
            gate_label = "通过硬门槛" if gate.get("passed") else "Research / Shadow"
            cards.append(f"""
              <article class="model-card" data-family="{html.escape(str(family))}" data-status="{html.escape(str(model.get('status')))}">
                <header><div><small>{html.escape(str(family))}</small><h3>{html.escape(str(title))}</h3></div><span class="gate {'pass' if gate.get('passed') else 'hold'}">{gate_label}</span></header>
                <div class="metric-grid">{rows_html}</div>
                {f'<div class="chart-row">{visuals}</div>' if visuals else ''}
                {task_detail}
                <p>{html.escape(str(model.get('reason') or model.get('rationale') or model.get('note') or 'Evidence persisted locally.'))}</p>
                {f"<details><summary>阻断项</summary><pre>{html.escape(json.dumps(gate.get('failedChecks') or model.get('blockingIssues') or [], ensure_ascii=False, indent=2))}</pre></details>" if not gate.get("passed") else ""}
              </article>
            """)
        market_sections.append(f"""
          <section class="market-section" data-market="{market['market']}">
            <div class="market-head"><div><span>{market['market']}</span><h2>{html.escape(str(market.get('modelVersion') or '无市场级注册模型'))}</h2></div>
              <div class="market-stats"><b>{dataset.get('rawRows', 0)}</b><small>训练行</small><b>{dataset.get('symbolCount', 0)}</b><small>股票</small><b>{dataset.get('dateCount', 0)}</b><small>日期</small></div>
            </div>
            <p class="market-policy">活跃特征 {dataset.get('activeFeatureCount', 0)} 个 · 已训练周期 {html.escape(str(market.get('trainedHorizons') or []))} · 暂缓周期 {html.escape(str(market.get('withheldHorizons') or []))}。{html.escape(str(dataset.get('featurePolicy') or ''))}</p>
            <div class="model-list">{''.join(cards)}</div>
          </section>
        """)
    job = evidence["jobReliability"]
    document = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{html.escape(evidence['title'])}</title>
<style>
:root{{--bg:#090c10;--surface:#11161c;--surface2:#182129;--ink:#f3f7f8;--muted:#98a4a9;--gold:#d6b46e;--green:#72d69a;--red:#f18492;--line:#2b343a}}
*{{box-sizing:border-box}}body{{margin:0;background:var(--bg);color:var(--ink);font:15px/1.55 -apple-system,BlinkMacSystemFont,"PingFang SC","Microsoft YaHei",sans-serif}}
main{{width:min(1240px,94vw);margin:auto;padding:48px 0 80px}}.hero{{padding:38px;border:1px solid var(--line);border-radius:18px;background:linear-gradient(140deg,#161c22,#0d1115)}}
.hero small,.model-card small{{color:var(--gold);letter-spacing:.12em;text-transform:uppercase}}h1{{font-size:34px;margin:8px 0}}h2,h3{{margin:0}}.summary{{display:grid;grid-template-columns:repeat(4,1fr);gap:10px;margin-top:24px}}
.summary div,.metric-grid div{{background:rgba(255,255,255,.035);border:1px solid rgba(255,255,255,.06);border-radius:10px;padding:12px}}.summary b{{font-size:22px;display:block}}
.warning{{margin:20px 0 34px;padding:16px 18px;border-left:3px solid var(--red);background:rgba(241,132,146,.07);border-radius:8px}}
.market-section{{margin-top:42px}}.market-head{{display:flex;justify-content:space-between;align-items:end;margin-bottom:14px}}.market-head>div>span{{color:var(--gold)}}
.market-policy{{margin:0 0 14px;color:var(--muted)}}
.market-stats{{display:grid;grid-template-columns:repeat(3,auto auto);gap:5px 10px;align-items:baseline}}.market-stats small{{color:var(--muted)}}.model-list{{display:grid;gap:12px}}
.model-card{{padding:20px;border-radius:14px;border:1px solid var(--line);background:var(--surface)}}.model-card header{{display:flex;justify-content:space-between;gap:14px;align-items:start}}
.gate{{font-size:12px;border-radius:999px;padding:5px 9px}}.gate.pass{{color:var(--green);background:rgba(114,214,154,.09)}}.gate.hold{{color:var(--gold);background:rgba(214,180,110,.09)}}
.metric-grid{{display:grid;grid-template-columns:repeat(4,1fr);gap:8px;margin:16px 0}}.metric-grid div{{display:flex;justify-content:space-between;gap:8px;padding:9px}}.metric-grid span{{color:var(--muted)}}
.filters{{position:sticky;top:0;z-index:3;display:flex;gap:8px;align-items:center;margin:0 0 30px;padding:12px 0;background:rgba(9,12,16,.94);backdrop-filter:blur(14px)}}
.filters button,.filters select{{min-height:36px;border:1px solid var(--line);border-radius:9px;background:var(--surface);color:var(--ink);padding:0 12px}}.filters button.active{{border-color:var(--gold);color:var(--gold)}}
.chart-row{{display:grid;grid-template-columns:1.4fr .6fr;gap:10px;margin:12px 0}}.evidence-chart{{margin:0;padding:12px;border:1px solid rgba(255,255,255,.06);border-radius:10px;background:#0d1216}}
.evidence-chart figcaption{{margin-bottom:8px;color:var(--muted);font-size:12px}}.evidence-chart svg{{display:block;width:100%;height:190px}}.evidence-chart line.ideal{{stroke:#536068;stroke-dasharray:5 5}}
.evidence-chart polyline.actual{{fill:none;stroke:var(--gold);stroke-width:3}}.evidence-chart circle{{fill:var(--gold)}}.evidence-chart text{{fill:var(--muted);font-size:11px}}
.confusion-grid{{display:grid;grid-template-columns:1fr 1fr;gap:8px}}.confusion-grid span{{display:grid;place-items:center;min-height:76px;border-radius:8px;background:rgba(214,180,110,.08)}}.confusion-grid small{{color:var(--muted)}}.confusion-grid b{{font-size:20px}}
.task-evidence{{margin:12px 0;padding:10px 12px;border:1px solid rgba(255,255,255,.06);border-radius:10px}}.task-evidence table{{width:100%;margin-top:10px;border-collapse:collapse;font-size:13px}}.task-evidence th,.task-evidence td{{padding:8px;border-bottom:1px solid rgba(255,255,255,.06);text-align:left}}.task-evidence th{{color:var(--gold)}}
p{{color:#c8d0d3}}details{{color:var(--muted)}}pre{{white-space:pre-wrap;color:#d7dddf}}[hidden]{{display:none!important}}@media(max-width:720px){{.summary,.metric-grid,.chart-row{{grid-template-columns:1fr 1fr}}.market-head{{align-items:start;gap:16px}}.market-stats{{grid-template-columns:auto auto}}.filters{{overflow:auto}}}}
</style></head><body><main>
<section class="hero"><small>MODEL GOVERNANCE REPORT</small><h1>{html.escape(evidence['title'])}</h1><p>生成时间 {html.escape(evidence['generatedAt'])} · 报告编号 {html.escape(evidence['reportId'])}</p>
<div class="summary"><div><b>{'是' if evidence['productionReady'] else '否'}</b><span>生产就绪</span></div><div><b>{job['total']}</b><span>后台任务</span></div><div><b>{_fmt((job.get('currentRuntime') or {}).get('terminalFailureRatePct', job.get('terminalFailureRatePct')), suffix='%')}</b><span>升级后失败率</span></div><div><b>{len(evidence['blockers'])}</b><span>当前阻断项</span></div></div></section>
<div class="warning">{html.escape(evidence['honestBoundary'])}</div>
<nav class="filters" aria-label="报告筛选">
  <button type="button" data-market-filter="ALL" class="active">全部市场</button>
  {''.join(f'<button type="button" data-market-filter="{market}">{market}</button>' for market in evidence['scope'])}
  <select id="familyFilter"><option value="ALL">全部模型族</option><option value="market_multitask">市场多任务</option><option value="factor_research">因子研究</option><option value="intraday">分钟模型</option><option value="paper_agent">Paper Agent</option><option value="ai_supervisor">AI 监工</option></select>
</nav>
{''.join(market_sections)}
</main><script>
const marketButtons=[...document.querySelectorAll("[data-market-filter]")];
const familyFilter=document.querySelector("#familyFilter");
let activeMarket="ALL";
function applyFilters(){{
  document.querySelectorAll(".market-section").forEach(section=>{{
    const marketVisible=activeMarket==="ALL"||section.dataset.market===activeMarket;
    let visibleCards=0;
    section.querySelectorAll(".model-card").forEach(card=>{{
      const familyVisible=familyFilter.value==="ALL"||card.dataset.family===familyFilter.value;
      card.hidden=!(marketVisible&&familyVisible);
      if(!card.hidden) visibleCards+=1;
    }});
    section.hidden=!marketVisible||visibleCards===0;
  }});
}}
marketButtons.forEach(button=>button.addEventListener("click",()=>{{
  activeMarket=button.dataset.marketFilter;
  marketButtons.forEach(row=>row.classList.toggle("active",row===button));
  applyFilters();
}}));
familyFilter.addEventListener("change",applyFilters);
</script></body></html>"""
    target.write_text(document, "utf-8")


DOCX_FONT_FAMILY = "Hiragino Sans GB"


def _set_docx_font(target: Any, font_name: str = DOCX_FONT_FAMILY) -> None:
    from docx.oxml.ns import qn

    target.font.name = font_name
    fonts = target._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)


def _set_cell_text(cell: Any, text: str, *, bold: bool = False, color: str = "25313A") -> None:
    from docx.shared import Pt, RGBColor

    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    _set_docx_font(run)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(9)


def _repeat_table_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def _set_table_widths(table: Any, widths: list[float]) -> None:
    from docx.shared import Inches

    table.autofit = False
    for index, width in enumerate(widths):
        if index >= len(table.columns):
            break
        table.columns[index].width = Inches(width)
        for cell in table.columns[index].cells:
            cell.width = Inches(width)


def _render_evidence_chart(metrics: dict[str, Any], target: Path) -> bool:
    points = metrics.get("reliabilityCurve") or []
    matrix = metrics.get("confusionMatrix") or {}
    if not points or not matrix:
        return False
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return False
    width, height = 1200, 420
    image = Image.new("RGB", (width, height), "#F6F7F7")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    gold, ink, muted, line = "#9A7536", "#25313A", "#66757D", "#D9DEDF"
    chart_left, chart_top, chart_right, chart_bottom = 70, 56, 720, 360
    draw.text((chart_left, 18), "OOF RELIABILITY CURVE", fill=ink, font=font)
    draw.line((chart_left, chart_bottom, chart_right, chart_top), fill=line, width=3)
    coords = []
    for row in points:
        x = chart_left + clamp(row.get("predictedPct"), 0, 100) / 100 * (chart_right - chart_left)
        y = chart_bottom - clamp(row.get("actualPct"), 0, 100) / 100 * (chart_bottom - chart_top)
        coords.append((x, y))
    if len(coords) > 1:
        draw.line(coords, fill=gold, width=5)
    for x, y in coords:
        draw.ellipse((x - 6, y - 6, x + 6, y + 6), fill=gold)
    draw.text((chart_left, chart_bottom + 12), "Predicted probability 0% -> 100%", fill=muted, font=font)
    draw.text((810, 18), "CONFUSION MATRIX", fill=ink, font=font)
    boxes = (
        ("TN", matrix.get("tn"), (800, 58, 970, 190)),
        ("FP", matrix.get("fp"), (990, 58, 1160, 190)),
        ("FN", matrix.get("fn"), (800, 210, 970, 342)),
        ("TP", matrix.get("tp"), (990, 210, 1160, 342)),
    )
    for label, value, box in boxes:
        draw.rounded_rectangle(box, radius=12, fill="#EEE8DC", outline="#D5C39C", width=2)
        draw.text((box[0] + 18, box[1] + 18), label, fill=muted, font=font)
        draw.text((box[0] + 18, box[1] + 65), str(int(number(value))), fill=ink, font=font)
    target.parent.mkdir(parents=True, exist_ok=True)
    image.save(target, "PNG")
    return True


def render_docx(evidence: dict[str, Any], target: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("python-docx is required to generate Word reports.") from exc

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.82)
    section.left_margin = section.right_margin = Inches(0.86)
    styles = doc.styles
    _set_docx_font(styles["Normal"])
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor(37, 49, 58)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name, size, color in (("Heading 1", 16, "715424"), ("Heading 2", 13, "715424"), ("Heading 3", 11, "25313A")):
        style = styles[style_name]
        _set_docx_font(style)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "GLOBAL QUANT WATCH  /  MODEL GOVERNANCE"
    _set_docx_font(header.runs[0])
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("715424")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Confidential local research evidence  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("MODEL TRAINING & PRODUCTION GATE REPORT")
    run.bold = True
    _set_docx_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("A17E3E")
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("Global Quant Watch\nAll-model training and production-gate report")
    title_run.bold = True
    _set_docx_font(title_run)
    title_run.font.size = Pt(25)
    title_run.font.color.rgb = RGBColor.from_string("172127")
    meta = doc.add_paragraph(f"Report ID: {evidence['reportId']}    Generated: {evidence['generatedAt']}")
    meta.runs[0].font.color.rgb = RGBColor.from_string("66757D")
    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(10)
    lead.paragraph_format.space_after = Pt(12)
    lead_text = (
        "Conclusion: deterministic production gates passed; explicit Shadow, Paper, and Production promotion is still required."
        if evidence["productionReady"]
        else "Conclusion: current evidence does not meet production gates; models remain Research or Shadow only."
    )
    lead_run = lead.add_run(lead_text)
    lead_run.bold = True
    lead_run.font.color.rgb = RGBColor.from_string("9B1C1C")

    doc.add_heading("1. Executive summary", level=1)
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.autofit = False
    _repeat_table_header(summary_table.rows[0])
    for cell in summary_table.rows[0].cells:
        cell.width = Inches(1.56)
    values = [
        ("Production ready", "Yes" if evidence["productionReady"] else "No"),
        ("Background jobs", evidence["jobReliability"]["total"]),
        ("Runtime failure", _fmt((evidence["jobReliability"].get("currentRuntime") or {}).get("terminalFailureRatePct", evidence["jobReliability"].get("terminalFailureRatePct")), suffix="%")),
        ("Blockers", len(evidence["blockers"])),
    ]
    for index, (label, value) in enumerate(values):
        _set_cell_text(summary_table.cell(0, index), label, bold=True, color="715424")
        _set_cell_text(summary_table.cell(1, index), value, bold=True)
    doc.add_paragraph(evidence["honestBoundary"])

    doc.add_heading("2. Market and model evidence", level=1)
    for market in evidence["markets"]:
        dataset = market.get("dataset") or {}
        doc.add_heading(f"{market['market']} - {market.get('modelVersion') or 'no registered market model'}", level=2)
        doc.add_paragraph(
            f"Raw training rows {int(number(dataset.get('rawRows')))}; symbols {int(number(dataset.get('symbolCount')))}; "
            f"dates {int(number(dataset.get('dateCount')))}; PIT event coverage {_fmt(dataset.get('pointInTimeCoveragePct'), suffix='%')}; "
            f"verified corporate-action history {_fmt(dataset.get('corporateActionCoveragePct'), suffix='%')}; "
            f"adjusted-price coverage {_fmt(dataset.get('adjustedPriceCoveragePct'), suffix='%')}."
        )
        doc.add_paragraph(
            f"Active features {int(number(dataset.get('activeFeatureCount')))}; "
            f"trained horizons {market.get('trainedHorizons') or []}; withheld horizons {market.get('withheldHorizons') or []}. "
            f"{dataset.get('featurePolicy') or ''}"
        )
        model_table = doc.add_table(rows=1, cols=7)
        model_table.style = "Table Grid"
        _repeat_table_header(model_table.rows[0])
        headers = ("Model", "Task", "Status", "Samples", "Acc.", "Prec.", "F1")
        for index, label in enumerate(headers):
            _set_cell_text(model_table.cell(0, index), label, bold=True, color="715424")
        factor_models = [model for model in market["models"] if model.get("family") == "factor_research"]
        report_models = [model for model in market["models"] if model.get("family") != "factor_research"]
        report_models.extend(factor_models[:12])
        for model in report_models:
            row = model_table.add_row().cells
            primary = next((item for item in model.get("classifiers") or [] if item["id"] == "directionProbability"), None) or next((item for item in model.get("classifiers") or [] if item["id"] == "ensembleProbability"), {})
            metrics = primary.get("metrics") or model.get("metrics") or {}
            raw_accuracy = metrics.get("accuracyPct")
            if isinstance(raw_accuracy, dict):
                raw_accuracy = raw_accuracy.get("value")
            raw_precision = metrics.get("precisionPct")
            if isinstance(raw_precision, dict):
                raw_precision = raw_precision.get("value")
            raw_f1 = metrics.get("f1Pct")
            if isinstance(raw_f1, dict):
                raw_f1 = raw_f1.get("value")
            values = (
                model.get("name") or model.get("modelId"),
                model.get("family"),
                model.get("status"),
                model.get("sampleAudit", {}).get("uniqueRows", model.get("sampleCount", 0)),
                _fmt(raw_accuracy, suffix="%"),
                _fmt(raw_precision, suffix="%"),
                _fmt(raw_f1, suffix="%"),
            )
            for index, value in enumerate(values):
                _set_cell_text(row[index], str(value))
        _set_table_widths(model_table, [1.30, 1.05, 0.80, 0.68, 0.68, 0.68, 0.50])
        if len(factor_models) > 12:
            doc.add_paragraph(
                f"The Word summary lists 12 of {len(factor_models)} factor artifacts; the HTML and JSON evidence retain the complete registry."
            )

        market_models = [model for model in market["models"] if model.get("family") == "market_multitask"]
        for model in market_models:
            doc.add_heading(f"{model.get('horizon')}-day market model", level=3)
            audit = model.get("sampleAudit") or {}
            doc.add_paragraph(
                f"Status: {model.get('status')}; OOF raw/deduplicated {audit.get('rawRows', 0)}/{audit.get('uniqueRows', 0)}; "
                f"independent dates {audit.get('independentDates', 0)}; duplicates {audit.get('duplicateRows', 0)}; "
                f"cross-market contamination {audit.get('crossMarketRows', 0)}."
            )
            failures = model.get("hardGate", {}).get("failedChecks") or []
            doc.add_paragraph("Hard gate: " + ("passed" if not failures else "failed: " + ", ".join(failures)))
            task_table = doc.add_table(rows=1, cols=6)
            task_table.style = "Table Grid"
            _repeat_table_header(task_table.rows[0])
            for index, label in enumerate(("Model/output", "Task", "Samples", "Accuracy/IC", "F1/MAE", "Brier/Lift")):
                _set_cell_text(task_table.cell(0, index), label, bold=True, color="715424")
            for classifier in model.get("classifiers") or []:
                values = classifier.get("metrics") or {}
                row = task_table.add_row().cells
                task_values = (
                    classifier.get("name"),
                    "Classification",
                    _fmt(values.get("samples"), 0),
                    _fmt(values.get("accuracyPct"), suffix="%"),
                    _fmt(values.get("f1Pct"), suffix="%"),
                    _fmt(values.get("brierSkillScore"), 3),
                )
                for index, value in enumerate(task_values):
                    _set_cell_text(row[index], value)
            for name, task, values, columns in (
                ("Cross-sectional ranking", "Ranking", model.get("ranking") or {}, ("dateCount", "rankIc", "precisionAt10Pct", "topDecileLiftPct")),
                ("Median return", "Regression", model.get("regression") or {}, ("samples", "directionAccuracyPct", "mae", "r2")),
                ("Return quantiles", "Quantile", model.get("quantiles") or {}, ("samples", "intervalCoveragePct", "p50Pinball", "meanIntervalWidthPct")),
            ):
                row = task_table.add_row().cells
                task_values = (name, task, *(_fmt(values.get(key)) for key in columns))
                for index, value in enumerate(task_values):
                    _set_cell_text(row[index], value)
            _set_table_widths(task_table, [1.55, 0.85, 0.75, 1.00, 0.90, 1.00])
            importance = model.get("featureImportance") or []
            if importance:
                top_features = importance[:12]
                doc.add_paragraph(
                    "Top feature importance: " + "; ".join(
                        f"{row.get('name', row.get('feature', 'unknown'))}={_fmt(row.get('importance', row.get('value')), 3)}"
                        for row in top_features
                    )
                )
            else:
                doc.add_paragraph("Feature importance: the current registered artifact does not contain verifiable importance evidence.")
            primary = next((row for row in model.get("classifiers") or [] if row["id"] == "directionProbability"), None) or next((row for row in model.get("classifiers") or [] if row["id"] == "ensembleProbability"), {})
            primary_metrics = primary.get("metrics") or {}
            chart_path = target.parent / "assets" / f"{evidence['reportId']}-{market['market']}-{model.get('horizon')}d.png"
            if _render_evidence_chart(primary_metrics, chart_path):
                doc.add_picture(str(chart_path), width=Inches(6.7))
                caption = doc.add_paragraph("Figure: strict OOF reliability curve and confusion matrix")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].italic = True
                caption.runs[0].font.size = Pt(8.5)
            else:
                doc.add_paragraph("OOF chart unavailable because reliability-curve or confusion-matrix evidence is insufficient.")

    doc.add_heading("3. Paper Agents and AI supervisors", level=1)
    for market in evidence["markets"]:
        agents = [model for model in market["models"] if model.get("family") == "paper_agent"]
        reviewers = [model for model in market["models"] if model.get("family") == "ai_supervisor"]
        doc.add_heading(market["market"], level=2)
        for agent in agents:
            values = agent.get("metrics") or {}
            doc.add_paragraph(
                f"{agent.get('name') or agent.get('modelId')}: net return after costs {_fmt(values.get('netReturnPct'), suffix='%')}; "
                f"trades {values.get('tradeCount', 0)}; win rate {_fmt(values.get('winRatePct'), suffix='%')}."
            )
        for reviewer in reviewers:
            doc.add_paragraph(
                f"{reviewer.get('provider')}: verdict={reviewer.get('verdict')}; available={reviewer.get('available')}."
            )

    doc.add_heading("4. Production gates and rework", level=1)
    gate_table = doc.add_table(rows=1, cols=3)
    gate_table.style = "Table Grid"
    _repeat_table_header(gate_table.rows[0])
    for index, label in enumerate(("Gate", "Requirement", "Purpose")):
        _set_cell_text(gate_table.cell(0, index), label, bold=True, color="715424")
    requirements = (
        ("Samples", "50,000 train rows, 1,000 OOF rows, 120 test dates", "Prevent small-sample inflation"),
        ("Events", "500 target and stop events; at least 4 of 5 positive folds", "Cover distinct market regimes"),
        ("Probability", "Brier Skill > 0, ECE <= 5%, slope 0.8-1.2", "Keep probabilities interpretable"),
        ("Quality", "Zero contamination, duplicates, and leakage; PSI <= 0.40", "Block leakage and drift"),
        ("Return", "Positive Top-K Lift and net EV", "Require economic value after costs"),
        ("Supervision", "At least two available AI supervisors accept", "Review only after deterministic gates"),
    )
    for requirement in requirements:
        cells = gate_table.add_row().cells
        for index, value in enumerate(requirement):
            _set_cell_text(cells[index], value)
    _set_table_widths(gate_table, [0.90, 3.25, 2.10])

    doc.add_heading("5. Current remediation queue", level=1)
    for action in evidence["recommendedActions"]:
        doc.add_paragraph(action, style="List Bullet")
    doc.add_paragraph(
        "Paper Agent returns, rule heads, AI reviewer opinions, and duplicate predictions are never relabeled as classification accuracy. "
        "Missing metrics remain unavailable until a strict OOF artifact can support them."
    )
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)


def generate_model_report(
    *,
    root: str | Path | None = None,
    markets: list[str] | None = None,
    output_dir: str | Path | None = None,
) -> dict[str, Any]:
    app_root = Path(root or Path(__file__).resolve().parents[1]).resolve()
    evidence = build_report_evidence(app_root, markets)
    base = Path(output_dir).resolve() if output_dir else app_root / ".cache" / "model-reports" / evidence["reportId"]
    base.mkdir(parents=True, exist_ok=True)
    json_path = base / "evidence.json"
    html_path = base / "report.html"
    docx_cache_path = base / f"{evidence['reportId']}.docx"
    json_path.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), "utf-8")
    render_html(evidence, html_path)
    render_docx(evidence, docx_cache_path)
    permanent_dir = app_root / "reports" / "model-training"
    permanent_dir.mkdir(parents=True, exist_ok=True)
    permanent_path = permanent_dir / docx_cache_path.name
    permanent_path.write_bytes(docx_cache_path.read_bytes())
    index_path = app_root / ".cache" / "model-reports" / "index.json"
    index = read_json(index_path, {"reports": []})
    entry = {
        "reportId": evidence["reportId"],
        "generatedAt": evidence["generatedAt"],
        "scope": evidence["scope"],
        "productionReady": evidence["productionReady"],
        "basePath": str(base),
        "jsonPath": str(json_path),
        "htmlPath": str(html_path),
        "docxPath": str(docx_cache_path),
        "permanentDocxPath": str(permanent_path),
    }
    reports = [entry, *[row for row in index.get("reports") or [] if row.get("reportId") != entry["reportId"]]][:100]
    index_path.parent.mkdir(parents=True, exist_ok=True)
    index_path.write_text(json.dumps({"updatedAt": iso_now(), "reports": reports}, ensure_ascii=False, indent=2), "utf-8")
    return {**entry, "evidence": evidence}
